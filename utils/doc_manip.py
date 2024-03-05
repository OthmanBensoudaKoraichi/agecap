import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
#from docx2pdf import convert
import os


def insert_into_word_doc(doc_path, list_members, df_premiums,id_devis):
    # Define the doc
    doc = Document(doc_path)

    # Modify df_offer
    doc.tables[0].cell(1, 2).text = id_devis
    doc.tables[0].cell(1, 3).text = datetime.datetime.today().strftime("%d-%m-%Y")
    doc.tables[0].cell(1, 4).text = (
        datetime.datetime.today() + datetime.timedelta(days=30)
    ).strftime("%d-%m-%Y")

    # Center align the text
    for cell_number in range(2, 5):
        for paragraph in doc.tables[0].cell(1, cell_number).paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Iterate over family members to format and insert data into the second table
    for i, member in enumerate(list_members):
        # Assuming the date of birth is at index 2 in the member tuple
        year, month, day = str(member[2]).split("-")
        date_reformatted = f"{day[:2]}-{month}-{year}"

        # Assuming the relation type is at index 3 in the member tuple
        relation = member[3]

        # Insert the date into the specified cell (column 2)
        cell_date = doc.tables[1].cell(i + 1, 2)
        cell_date.text = date_reformatted
        cell_date.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Insert the relation into the specified cell (column 1)
        cell_relation = doc.tables[1].cell(i + 1, 1)

        # Clear any existing paragraphs in the cell by removing all but the first paragraph
        for para in cell_relation.paragraphs[1:]:
            p = para._element
            p.getparent().remove(p)

        # Modify the text of the first (and now only) paragraph in the cell
        cell_relation.paragraphs[0].text = relation
        cell_relation.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Modify df_premiums
    for i in range(df_premiums.shape[0]):  # df.shape[0] gives the number of rows
        for j in range(df_premiums.shape[1]):  # df.shape[1] gives the number of columns
            cell_value = str(df_premiums.iloc[i, j])
            cell = doc.tables[2].cell(i + 2, j + 1)
            cell.text = cell_value

            # Center align the text in the cell
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Use tempfile to create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        docx_path = tmp.name

    # Convert the docx to pdf
    #pdf_path = docx_path.replace(".docx", ".pdf")
    #convert(docx_path, pdf_path)

    # Make sure to remove the temporary docx file if you no longer need it
    #os.unlink(docx_path)

    #return pdf_path

    return docx_path

def insert_into_html(html_code, list_members, df_premiums, id_devis):
    # Format today's date and the date 30 days from now
    today_date = datetime.datetime.today().strftime("%d-%m-%Y")
    valid_until_date = (datetime.datetime.today() + datetime.timedelta(days=30)).strftime("%d-%m-%Y")

    # Correctly replace the placeholder for full name
    html_code = html_code.replace('id="full_name">Placeholder</span>', f'id="full_name">{list_members[0][0].capitalize() + " " + list_members[0][1].capitalize()}</span>')
    html_code = html_code.replace('<td id="devis-number">Placeholder</td>', f'<td id="devis-number">{id_devis}</td>')
    html_code = html_code.replace('<td id="issued-date">Placeholder</td>', f'<td id="issued-date">{today_date}</td>')
    html_code = html_code.replace('<td id="valid-until-date">Placeholder</td>', f'<td id="valid-until-date">{valid_until_date}</td>')

    # Insert family members' information
    members_html = ''
    for i, member in enumerate(list_members):
        year, month, day = str(member[2]).split("-")
        date_reformatted = f"{day[:2]}-{month}-{year}"
        relation = member[3]
        members_html += f'<tr><td>{relation}</td><td>{date_reformatted}</td></tr>'

    html_code = html_code.replace('<!-- Placeholder for Members Information -->', members_html)

    premiums_html = ''
    for i in range(df_premiums.shape[0]):
        row_html = f'<tr><td class="dark-blue"><strong>{labels[i]}</strong></td>'
        for j in range(df_premiums.shape[1]):
            cell_value = df_premiums.iloc[i, j]
            # Apply dark blue to specific percentages
            if cell_value in ["80%", "90%"]:  # Assuming you want to add "85%" if it exists
                row_html += f'<td class="dark-blue">{cell_value}</td>'
            else:
                row_html += f'<td>{cell_value}</td>'
        row_html += '</tr>'
        premiums_html += row_html

    html_code = html_code.replace('<!-- Placeholder for Premiums Information -->', premiums_html)

    return html_code